# Install openpyxl and docx modules on your system

from tkinter import *
import tkinter as tk
from tkinter import PhotoImage
from openpyxl import load_workbook
from datetime import date, timedelta
from tkinter import messagebox
import json
import re
from docx import Document
import os

# Create the main window
root = tk.Tk()
root.title("Movie Ticket Booking System")

# Get the screen width and height
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# Function to exit the application
def exit_application():
    root.destroy()

# Function to open a new window for movie selection
def open_new_window():
    new_window = tk.Toplevel(root)
    new_window.title("Movie selection")
    
    close_button = tk.Button(new_window, text="Close", command=new_window.destroy)
    close_button.place(relx=0.05, rely=0.9)
    new_window.geometry(f"{screen_width}x{screen_height}")
    
    # Load the Excel file containing cities
    wb1= load_workbook('Cities.xlsx')
    ws1= wb1.active

    #Load the Excel file containing movies
    wb2=load_workbook('Movies.xlsx')
    ws2= wb2.active
    
    # Store the city names in a list
    cities = []
    for row in ws1.iter_rows(min_row=1, max_col=1, values_only=True):
        if row[0] not in cities:
            cities.append(row[0])

    # Label for city selection
    label = tk.Label(new_window, text="Select a city from the list:")
    label.pack(pady=10)
    
    # Create the dropdown menu for cities
    city_var = StringVar(new_window)
    city_var.set(cities[0]) # set the default option
    city_dropdown = OptionMenu(new_window, city_var, *cities)
    city_dropdown.pack()

    # Create the dropdown menu for areas
    def update_areas(*args):
        areas = []
        for row in ws1.iter_rows(min_row=1, values_only=True):
            if row[0] == city_var.get():
                areas.extend(row[1:])
        area_var.set(areas[0])
        area_dropdown['menu'].delete(0, 'end')
        for area in areas:
            area_dropdown['menu'].add_command(label=area, command=lambda value=area: area_var.set(value))

    label2= tk.Label(new_window, text="Select the area from the list:")
    label2.pack(pady=10)
    area_var = StringVar(new_window)
    area_dropdown = OptionMenu(new_window, area_var, "")
    city_var.trace('w', update_areas) # update areas when the city is changed
    update_areas() # set the initial areas
    area_dropdown.pack()

    # Create the button to print the selected city and area
    def Select_city_and_area():
        result_label1.config(text=f"Selected City: {city_var.get()}")
        result_label2.config(text=f"Selected Area: {area_var.get()}")

    button1= tk.Button(new_window, text="Select the city and area ", command=Select_city_and_area)
    button1.pack()

    result_label1= tk.Label(new_window, text="")
    result_label1.pack()

    result_label2= tk.Label(new_window, text="")
    result_label2.pack()

    # Label for date selection    
    label4 = tk.Label(new_window, text="Select the date :")
    label4.pack(pady=10)

    # Get the current date
    current_date = date.today()

    # Format the date as dd-mm-yy
    curr_date = current_date.strftime("%d-%m-%y")

    # Calculate the five dates next to the current date
    date1 = current_date + timedelta(days=1)
    date_1 = date1.strftime("%d-%m-%y")

    date2 = current_date + timedelta(days=2)
    date_2 = date2.strftime("%d-%m-%y")

    date3 = current_date + timedelta(days=3)
    date_3 = date3.strftime("%d-%m-%y")

    date4 = current_date + timedelta(days=4)
    date_4 = date4.strftime("%d-%m-%y")

    # Create a label to display the current date
    current_date_label = tk.Label(new_window, text="Current Date: " + str(curr_date))
    current_date_label.pack()

    # Create radio buttons for the five dates
    selected_date = tk.StringVar()

    date1_radio = tk.Radiobutton(new_window, text=str(curr_date), variable=selected_date, value=str(curr_date))
    date1_radio.pack()

    date2_radio = tk.Radiobutton(new_window, text=str(date_1), variable=selected_date, value=str(date_1))
    date2_radio.pack()

    date3_radio = tk.Radiobutton(new_window, text=str(date_2), variable=selected_date, value=str(date_2))
    date3_radio.pack()

    date4_radio = tk.Radiobutton(new_window, text=str(date_3), variable=selected_date, value=str(date_3))
    date4_radio.pack()

    date5_radio = tk.Radiobutton(new_window, text=str(date_4), variable=selected_date, value=str(date_4))
    date5_radio.pack()

    # Set the default option (Option 1 in this case)
    selected_date.set(str(date_1))

    # Function to get the selected date
    def get_selected_date():
        result_label4.config(text=f"Selected Date: {selected_date.get()}")

    # Create a button to get the selected date
    get_date_button = tk.Button(new_window, text="Get Selected Date", command=get_selected_date)
    get_date_button.pack()

    # Label to display the selected date
    result_label4= tk.Label(new_window, text="")
    result_label4.pack()

    # Create the dropdown menu for movies
    def update_movies(*args):
        movies = []
        for row in ws2.iter_rows(min_row=1, values_only=True):
            if row[0] == city_var.get():
                movies.extend(row[1:])
        movie_var.set(movies[0])
        movie_dropdown['menu'].delete(0, 'end')
        for movie in movies:
            movie_dropdown['menu'].add_command(label=movie, command=lambda value=movie: movie_var.set(value))

    # Label for movie selection
    label3 = tk.Label(new_window, text="Select the movie :")
    label3.pack(pady=10)
    movie_var = StringVar(new_window)
    movie_dropdown = OptionMenu(new_window, movie_var, "")
    city_var.trace('w', update_movies) # update movies when the city is changed
    update_movies() # set the initial movies
    movie_dropdown.pack()
    
    # Create the button to print the selected values
    def Select_movie():
        result_label3.config(text=f"Selected Movie: {movie_var.get()}")

    button3 = Button(new_window, text="Select the movie", command=Select_movie)
    button3.pack()

    result_label3= tk.Label(new_window, text="")
    result_label3.pack()

    # Load the Excel sheet containg timings of show
    wb3=load_workbook(f"{city_var.get()}.xlsx")
    ws3= wb3.active

    # Create the dropdown menu for timings
    def update_timings(*args):
        timings = []
        for row in ws3.iter_rows(min_row=1, values_only=True):
            if row[0] == movie_var.get():
                timings.extend(row[1:])
        time_var.set(timings[0])
        time_dropdown['menu'].delete(0, 'end')
        for time in timings:
            time_dropdown['menu'].add_command(label=time, command=lambda value=time: time_var.set(value))

    # Label for timings
    label3 = tk.Label(new_window, text="Select the timings :")
    label3.pack(pady=10)
    time_var = StringVar(new_window)
    time_dropdown = OptionMenu(new_window, time_var, "")
    movie_var.trace('w', update_timings) # update timings when the movie is changed
    update_timings() # set the initial timings
    time_dropdown.pack()

    # Function to open a new window for seat selection
    def open_new_window_2():
        new_window_2= tk.Toplevel(root)
        new_window_2.title("Seat selection")
    
        close_button = tk.Button(new_window_2, text="Close", command=new_window_2.destroy)
        close_button.place(relx=0.05, rely=0.9)
        new_window_2.geometry(f"{screen_width}x{screen_height}")

        # Define the seat prices for each row
        seat_prices = [500] * 3 + [400] * 3 + [300] * (4)

        # Define a unique identifier for the show (e.g., show time or date)
        show_identifier = f"{movie_var.get()}_{city_var.get()}_{area_var.get()}_{selected_date.get()}_{time_var.get()}".replace(" ","_").replace(":","_") # Unique identifier for the show
        
        # File to store seat selection data
        selection_file = f"seat_selection_{show_identifier}.json"

        def save_selection():
            with open(selection_file, "w") as file:
                json.dump(seat_status, file)

        def load_selection():
            try:
                with open(selection_file, "r") as file:
                    return json.load(file)
            except FileNotFoundError:
                return [[0 for _ in range(num_cols)] for _ in range(num_rows)]

        def seat_click(row, col):
            if seat_status[row][col] == 0:
                seats[row][col]['bg'] = 'green'
                seat_status[row][col] = 1
                selected_seats.append((row, col))
                save_selection()
                update_seat_label()
            else:
                seats[row][col]['bg'] = 'grey'

        def update_seat_label():
            total_price = sum(seat_prices[row] for row, _ in selected_seats)
            seat_label.config(text=f"SCREEN IS HERE\nSelected Seats: {', '.join([f'{chr(65 + r)}{c + 1}' for r, c in selected_seats])}\nTotal Price: Rs{total_price}\nFirst three rows: Rs 500 per seat\n Next three rows: Rs 400 per seat\n Rest of the rows: Rs 300 per seat")


        # Define the number of rows and columns
        num_rows = 10
        num_cols = 25

        # Create a 2D list to store references to the seat labels/buttons
        seats = [[None for _ in range(num_cols)] for _ in range(num_rows)]

        # Create a 2D list to track seat availability
        seat_status = load_selection()

        # Create a list to keep track of selected seats
        selected_seats = []

        # Create labels/buttons for each seat
        for row in range(num_rows):
            for col in range(num_cols):
                seat_label = tk.Label(new_window_2, text=f'{chr(65 + row)}{col + 1}\nRs{seat_prices[row]}', width=5, height=3, relief='ridge', bg='white')
                if seat_status[row][col] == 1:
                    seat_label['bg'] = 'grey'
                seat_label.grid(row=row, column=col, padx=6, pady=5)
                seat_label.bind('<Button-1>', lambda event, row=row, col=col: seat_click(row, col))
                seats[row][col] = seat_label
        
        # Set row and column weights to make them expandable and centered
        for i in range(num_rows):
            new_window_2.grid_rowconfigure(i, weight=1)
        for i in range(num_cols):
             new_window_2.grid_columnconfigure(i, weight=1)

        # Create a label to display selected seats and the total price
        seat_label = tk.Label(new_window_2, text="SCREEN IS HERE\nSelected Seats: None\nTotal Price: Rs0\n First three rows: Rs 500 per seat\n Next three rows: Rs 400 per seat\n Rest of the rows: Rs 300 per seat")
        seat_label.grid(row=num_rows, columnspan=num_cols) 

        # Creating a payment window
        def open_payment_window():
            
            def ticket_details_window():

                # Create a new Tkinter window for ticket details
                details_window = tk.Toplevel(root)
                details_window.title("Ticket details")

                result_label6= tk.Label(details_window, text="")
                result_label6.pack()

                result_label7= tk.Label(details_window, text="")
                result_label7.pack()

                result_label8= tk.Label(details_window, text="")
                result_label8.pack()

                result_label9= tk.Label(details_window, text="")
                result_label9.pack()

                result_label10= tk.Label(details_window, text="")
                result_label10.pack()

                result_label11= tk.Label(details_window, text="")
                result_label11.pack()

                result_label6.config(text=f"Selected City: {city_var.get()}")
                result_label7.config(text=f"Selected Area: {area_var.get()}")
                result_label8.config(text=f"Selected Movie: {movie_var.get()}")
                result_label9.config(text=f"Selected Date: {selected_date.get()}")
                result_label10.config(text=f"Selected Time: {time_var.get()}")
                result_label11.config(text=f"Selected Seats: {[f'{chr(65 + r)}{c + 1}' for r, c in selected_seats]}")

            def create_ticket_document():
                
                # Define a unique identifier for the show (e.g., show time or date)
                show_identifier_2 = f"{movie_var.get()}_{city_var.get()}_{area_var.get()}_{selected_date.get()}_{time_var.get()}".replace(" ","_").replace(":","_") # Unique identifier for the show
        
                # File to store seat selection data
                filename = f"seat_selection_{show_identifier_2}.docx"
                
                if os.path.exists(filename):
                # If the file exists, open it and add new details
                    doc = Document(filename)
                # Add an empty line before every new entry
                    if doc.paragraphs:
                        doc.add_paragraph()

                else:
                # If the file doesn't exist, create a new document
                    doc = Document()
                    doc.add_heading('Ticket Details', level=1)
                    
                doc.add_paragraph(f"Selected City: {city_var.get()}")
                doc.add_paragraph(f"Selected Area: {area_var.get()}")
                doc.add_paragraph(f"Selected Movie: {movie_var.get()}")
                doc.add_paragraph(f"Selected Date: {selected_date.get()}")
                doc.add_paragraph(f"Selected Time: {time_var.get()}")
                doc.add_paragraph(f"Selected Seats: {', '.join([f'{chr(65 + r)}{c + 1}' for r, c in selected_seats])}")

                # Save the document with a unique filename
                doc.save(filename)

                return filename


            def process_payment():
                
                # Validate Card Number
                card = entry_card.get()
                if not card.isdigit() or len(card) != 12:
                    messagebox.showerror("Invalid Card Number", "Card Number is not valid.")
                    return
                
                # Validate Expiry Date
                expiry_date = entry_expiry.get()
                if not re.match(r"^(0[1-9]|1[0-2])\/\d{2}$", expiry_date):
                    messagebox.showerror("Invalid Expiry Date", "Expiry date must be in MM/YY format.")
                    return

                # Validate CVV
                cvv = entry_cvv.get()
                if not cvv.isdigit() or len(cvv) != 3:
                    messagebox.showerror("Invalid CVV", "CVV must be a 3-digit number.")
                    return
                
                # This function should implement the logic for processing the payment.
                # In a real application, you'd integrate with a payment gateway or service.
                # Here, we'll just display a confirmation message.

                messagebox.showinfo("Payment Received", "Payment has been successfully received.")
                ticket_details_window()
                create_ticket_document()

          
            # Create a new Tkinter window for payment
            payment_window = tk.Toplevel(root)
            payment_window.title("Payment Form")

            # Create labels and entry fields for payment information
            label_card = tk.Label(payment_window, text="Card Number (12 digits):")
            entry_card = tk.Entry(payment_window)

            label_expiry = tk.Label(payment_window, text="Expiration Date (MM/YY):")
            entry_expiry = tk.Entry(payment_window)

            label_cvv = tk.Label(payment_window, text="CVV (3 digits):")
            entry_cvv = tk.Entry(payment_window, show="*")  # Display * for security

            # Create a button to process payment
            payment_button = tk.Button(payment_window, text="Submit Payment", command=process_payment)

            # Pack widgets into the window
            label_card.pack()
            entry_card.pack()
            label_expiry.pack()
            entry_expiry.pack()
            label_cvv.pack()
            entry_cvv.pack()
            payment_button.pack()
        
        button5 = tk.Button(new_window_2, text="Pay", command=open_payment_window)
        button5.place(relx=0.9, rely=0.9)

    button4 = tk.Button(new_window, text="Next", command=open_new_window_2)
    button4.place(relx=0.9, rely=0.9) 

# Load the background image
bg_image = PhotoImage(file="theatre_background.png")

# Create a Label widget to display the background image
background_label = tk.Label(root, image=bg_image)
background_label.place(x=0, y=0, relwidth=1, relheight=1)

# Set the window size to fullscreen
root.geometry(f"{screen_width}x{screen_height}")

# Create two buttons
button1 = tk.Button(root, text="Exit", command=exit_application)
button2 = tk.Button(root, text="Next", command=open_new_window)

# Place the buttons in the leftmost and rightmost corners
button1.place(relx=0.05, rely=0.9)  # Adjust relx and rely as needed
button2.place(relx=0.9, rely=0.9)   # Adjust relx and rely as needed

root.mainloop()
